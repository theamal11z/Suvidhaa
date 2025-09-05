from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime
import base64
import io
import openai
import cloudinary
import cloudinary.uploader
import cloudinary.api
import PyPDF2
import docx
from PIL import Image
import json
import aiofiles

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Initialize services
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Configure Cloudinary
cloudinary.config(
    cloud_name=os.environ['CLOUDINARY_CLOUD_NAME'],
    api_key=os.environ['CLOUDINARY_API_KEY'],
    api_secret=os.environ['CLOUDINARY_API_SECRET']
)

# Configure NVIDIA API (using OpenAI client format)
nvidia_client = openai.OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key=os.environ['NVIDIA_API_KEY']
)

app = FastAPI(title="Suvidhaa API", description="Your Bridge to Transparent Governance")
api_router = APIRouter(prefix="/api")

# Models
class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    original_content: str
    summary_english: str
    summary_nepali: Optional[str] = None
    plain_language: str
    key_points: List[str]
    affected_groups: List[str]
    key_dates: List[str]
    responsible_offices: List[str]
    document_type: str
    file_url: Optional[str] = None
    file_base64: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    processed_at: Optional[datetime] = None

class DocumentCreate(BaseModel):
    title: str
    document_type: str

class Question(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_name: str
    email: str
    phone: Optional[str] = None
    question_text: str
    category: str
    related_document_id: Optional[str] = None
    evidence_urls: List[str] = []
    evidence_base64: List[str] = []
    government_office: str
    status: str = "submitted"  # submitted, routed, answered
    created_at: datetime = Field(default_factory=datetime.utcnow)
    response_text: Optional[str] = None
    response_at: Optional[datetime] = None

class QuestionCreate(BaseModel):
    user_name: str
    email: str
    phone: Optional[str] = None
    question_text: str
    category: str
    related_document_id: Optional[str] = None
    government_office: str

class Suggestion(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_name: str
    email: str
    suggestion_text: str
    category: str
    related_document_id: Optional[str] = None
    co_signatures: List[Dict[str, str]] = []
    status: str = "public"  # public, reviewed, implemented
    created_at: datetime = Field(default_factory=datetime.utcnow)
    sentiment_summary: Optional[str] = None

class SuggestionCreate(BaseModel):
    user_name: str
    email: str
    suggestion_text: str
    category: str
    related_document_id: Optional[str] = None

class Grievance(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_name: str
    email: str
    phone: str
    grievance_text: str
    category: str
    evidence_urls: List[str] = []
    evidence_base64: List[str] = []
    legal_references: List[str] = []
    affected_area: str
    government_office: str
    status: str = "filed"  # filed, under_review, resolved
    created_at: datetime = Field(default_factory=datetime.utcnow)
    resolution_text: Optional[str] = None
    resolved_at: Optional[datetime] = None

class GrievanceCreate(BaseModel):
    user_name: str
    email: str
    phone: str
    grievance_text: str
    category: str
    affected_area: str
    government_office: str

class Watchlist(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_email: str
    name: str
    keywords: List[str]
    categories: List[str]
    government_offices: List[str]
    notification_frequency: str = "daily"  # daily, weekly, monthly
    created_at: datetime = Field(default_factory=datetime.utcnow)
    last_notified: Optional[datetime] = None

class WatchlistCreate(BaseModel):
    user_email: str
    name: str
    keywords: List[str]
    categories: List[str]
    government_offices: List[str]
    notification_frequency: str = "daily"

# AI Processing Service
async def process_document_with_ai(content: str, title: str) -> Dict[str, Any]:
    """Process document content using NVIDIA AI"""
    try:
        # Create comprehensive prompt for document analysis
        prompt = f"""
        Analyze this government document and provide a structured response:
        
        Title: {title}
        Content: {content[:8000]}  # Limit content to avoid token limits
        
        Please provide:
        1. A concise summary in plain English (2-3 sentences)
        2. Key points (list 3-5 main points)
        3. Affected groups (who is impacted by this document)
        4. Important dates mentioned
        5. Responsible government offices
        6. Plain language explanation of complex terms
        
        Format as JSON with keys: summary, key_points, affected_groups, key_dates, responsible_offices, plain_language
        """
        
        response = nvidia_client.chat.completions.create(
            model=os.environ['NVIDIA_MODEL'],
            messages=[
                {"role": "system", "content": "You are an expert at simplifying government documents for citizens. Always respond in valid JSON format."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=1500
        )
        
        result = response.choices[0].message.content
        
        # Try to parse JSON response
        try:
            parsed_result = json.loads(result)
            return {
                "summary_english": parsed_result.get("summary", "Summary not available"),
                "key_points": parsed_result.get("key_points", []),
                "affected_groups": parsed_result.get("affected_groups", []),
                "key_dates": parsed_result.get("key_dates", []),
                "responsible_offices": parsed_result.get("responsible_offices", []),
                "plain_language": parsed_result.get("plain_language", "Plain language explanation not available")
            }
        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return {
                "summary_english": result[:500] + "..." if len(result) > 500 else result,
                "key_points": ["AI processing completed but structured data not available"],
                "affected_groups": ["General public"],
                "key_dates": [],
                "responsible_offices": ["To be determined"],
                "plain_language": "Please refer to the original document for detailed information."
            }
            
    except Exception as e:
        logging.error(f"AI processing error: {str(e)}")
        return {
            "summary_english": "AI processing temporarily unavailable. Document uploaded successfully.",
            "key_points": ["Document requires manual review"],
            "affected_groups": ["General public"],
            "key_dates": [],
            "responsible_offices": ["To be determined"],
            "plain_language": "AI analysis will be available shortly."
        }

# File processing utilities
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logging.error(f"PDF extraction error: {str(e)}")
        return "Error extracting text from PDF"

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logging.error(f"DOCX extraction error: {str(e)}")
        return "Error extracting text from DOCX"

# API Routes

@api_router.get("/")
async def root():
    return {"message": "Welcome to Suvidhaa API - Your Bridge to Transparent Governance"}

# UNDERSTAND Pillar - Document Processing
@api_router.post("/documents/upload", response_model=Document)
async def upload_document(
    file: UploadFile = File(...),
    title: str = Form(...),
    document_type: str = Form(...)
):
    try:
        # Read file content
        file_content = await file.read()
        
        # Extract text based on file type
        if file.content_type == "application/pdf":
            extracted_text = extract_text_from_pdf(file_content)
        elif file.content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            extracted_text = extract_text_from_docx(file_content)
        elif file.content_type.startswith("text/"):
            extracted_text = file_content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        # Upload to Cloudinary
        try:
            upload_result = cloudinary.uploader.upload(
                file_content,
                resource_type="raw",
                public_id=f"documents/{uuid.uuid4()}",
                use_filename=True
            )
            file_url = upload_result.get('secure_url')
        except Exception as e:
            logging.warning(f"Cloudinary upload failed: {str(e)}")
            file_url = None
        
        # Convert to base64 as fallback
        file_base64 = base64.b64encode(file_content).decode('utf-8')
        
        # Process with AI
        ai_result = await process_document_with_ai(extracted_text, title)
        
        # Create document object
        document = Document(
            title=title,
            original_content=extracted_text,
            document_type=document_type,
            file_url=file_url,
            file_base64=file_base64,
            processed_at=datetime.utcnow(),
            **ai_result
        )
        
        # Save to database
        await db.documents.insert_one(document.dict())
        
        return document
        
    except Exception as e:
        logging.error(f"Document upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@api_router.get("/documents", response_model=List[Document])
async def get_documents(skip: int = 0, limit: int = 20):
    documents = await db.documents.find().skip(skip).limit(limit).sort("created_at", -1).to_list(limit)
    return [Document(**doc) for doc in documents]

@api_router.get("/documents/{document_id}", response_model=Document)
async def get_document(document_id: str):
    document = await db.documents.find_one({"id": document_id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return Document(**document)

# ACT Pillar - Questions, Suggestions, Grievances
@api_router.post("/questions", response_model=Question)
async def submit_question(
    user_name: str = Form(...),
    email: str = Form(...),
    phone: Optional[str] = Form(None),
    question_text: str = Form(...),
    category: str = Form(...),
    related_document_id: Optional[str] = Form(None),
    government_office: str = Form(...),
    evidence_files: List[UploadFile] = File(default=[])
):
    try:
        evidence_urls = []
        evidence_base64 = []
        
        # Process evidence files
        for file in evidence_files:
            if file.filename:
                file_content = await file.read()
                
                # Upload to Cloudinary
                try:
                    upload_result = cloudinary.uploader.upload(
                        file_content,
                        resource_type="auto",
                        public_id=f"evidence/{uuid.uuid4()}",
                        use_filename=True
                    )
                    evidence_urls.append(upload_result.get('secure_url'))
                except Exception as e:
                    logging.warning(f"Evidence upload failed: {str(e)}")
                
                # Store as base64 fallback
                evidence_base64.append(base64.b64encode(file_content).decode('utf-8'))
        
        question = Question(
            user_name=user_name,
            email=email,
            phone=phone,
            question_text=question_text,
            category=category,
            related_document_id=related_document_id,
            government_office=government_office,
            evidence_urls=evidence_urls,
            evidence_base64=evidence_base64
        )
        
        await db.questions.insert_one(question.dict())
        return question
        
    except Exception as e:
        logging.error(f"Question submission error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Submission failed: {str(e)}")

@api_router.post("/suggestions", response_model=Suggestion)
async def submit_suggestion(suggestion_data: SuggestionCreate):
    try:
        suggestion = Suggestion(**suggestion_data.dict())
        await db.suggestions.insert_one(suggestion.dict())
        return suggestion
    except Exception as e:
        logging.error(f"Suggestion submission error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Submission failed: {str(e)}")

@api_router.post("/suggestions/{suggestion_id}/cosign")
async def cosign_suggestion(suggestion_id: str, signer_name: str = Form(...), signer_email: str = Form(...)):
    try:
        suggestion = await db.suggestions.find_one({"id": suggestion_id})
        if not suggestion:
            raise HTTPException(status_code=404, detail="Suggestion not found")
        
        # Add co-signature
        cosignature = {"name": signer_name, "email": signer_email, "signed_at": datetime.utcnow().isoformat()}
        
        await db.suggestions.update_one(
            {"id": suggestion_id},
            {"$push": {"co_signatures": cosignature}}
        )
        
        return {"message": "Co-signature added successfully"}
    except Exception as e:
        logging.error(f"Co-signature error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Co-signature failed: {str(e)}")

@api_router.post("/grievances", response_model=Grievance)
async def file_grievance(
    user_name: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    grievance_text: str = Form(...),
    category: str = Form(...),
    affected_area: str = Form(...),
    government_office: str = Form(...),
    evidence_files: List[UploadFile] = File(default=[])
):
    try:
        evidence_urls = []
        evidence_base64 = []
        
        # Process evidence files
        for file in evidence_files:
            if file.filename:
                file_content = await file.read()
                
                # Upload to Cloudinary
                try:
                    upload_result = cloudinary.uploader.upload(
                        file_content,
                        resource_type="auto",
                        public_id=f"grievance_evidence/{uuid.uuid4()}",
                        use_filename=True
                    )
                    evidence_urls.append(upload_result.get('secure_url'))
                except Exception as e:
                    logging.warning(f"Evidence upload failed: {str(e)}")
                
                # Store as base64 fallback
                evidence_base64.append(base64.b64encode(file_content).decode('utf-8'))
        
        grievance = Grievance(
            user_name=user_name,
            email=email,
            phone=phone,
            grievance_text=grievance_text,
            category=category,
            affected_area=affected_area,
            government_office=government_office,
            evidence_urls=evidence_urls,
            evidence_base64=evidence_base64
        )
        
        await db.grievances.insert_one(grievance.dict())
        return grievance
        
    except Exception as e:
        logging.error(f"Grievance filing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Filing failed: {str(e)}")

# TRACK Pillar - Watchlists and Dashboards
@api_router.post("/watchlists", response_model=Watchlist)
async def create_watchlist(watchlist_data: WatchlistCreate):
    try:
        watchlist = Watchlist(**watchlist_data.dict())
        await db.watchlists.insert_one(watchlist.dict())
        return watchlist
    except Exception as e:
        logging.error(f"Watchlist creation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Creation failed: {str(e)}")

@api_router.get("/watchlists", response_model=List[Watchlist])
async def get_watchlists(user_email: str):
    watchlists = await db.watchlists.find({"user_email": user_email}).to_list(100)
    return [Watchlist(**watchlist) for watchlist in watchlists]

@api_router.get("/dashboard/stats")
async def get_dashboard_stats():
    try:
        stats = {}
        
        # Document stats
        stats["total_documents"] = await db.documents.count_documents({})
        stats["documents_this_month"] = await db.documents.count_documents({
            "created_at": {"$gte": datetime.utcnow().replace(day=1, hour=0, minute=0, second=0, microsecond=0)}
        })
        
        # Question stats
        stats["total_questions"] = await db.questions.count_documents({})
        stats["answered_questions"] = await db.questions.count_documents({"status": "answered"})
        
        # Suggestion stats
        stats["total_suggestions"] = await db.suggestions.count_documents({})
        
        # Grievance stats
        stats["total_grievances"] = await db.grievances.count_documents({})
        stats["resolved_grievances"] = await db.grievances.count_documents({"status": "resolved"})
        
        return stats
    except Exception as e:
        logging.error(f"Dashboard stats error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Stats unavailable: {str(e)}")

@api_router.get("/submissions")
async def get_user_submissions(user_email: str):
    try:
        questions = await db.questions.find({"email": user_email}).sort("created_at", -1).to_list(50)
        suggestions = await db.suggestions.find({"email": user_email}).sort("created_at", -1).to_list(50)
        grievances = await db.grievances.find({"email": user_email}).sort("created_at", -1).to_list(50)
        
        return {
            "questions": [Question(**q) for q in questions],
            "suggestions": [Suggestion(**s) for s in suggestions],
            "grievances": [Grievance(**g) for g in grievances]
        }
    except Exception as e:
        logging.error(f"User submissions error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Submissions unavailable: {str(e)}")

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()